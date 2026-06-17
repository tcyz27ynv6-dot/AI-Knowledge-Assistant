from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import os, shutil
import chromadb, docx2txt, pandas as pd
from pypdf import PdfReader
from sentence_transformers import SentenceTransformer
from docx import Document
from openpyxl import Workbook
from CORE.rag import ask_rag

app = FastAPI(title="AI Knowledge Assistant")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])


class QuestionRequest(BaseModel):
    question: str


class ExportRequest(BaseModel):
    question: str
    answer: str
    sources: list


embedding_model = SentenceTransformer("BAAI/bge-small-en-v1.5")


@app.get("/")
def home():
    return {"message": "AI Knowledge Assistant API Running"}


@app.post("/ask")
def ask_question(request: QuestionRequest):
    return ask_rag(request.question)


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    os.makedirs("UPLOADS", exist_ok=True)
    file_path = os.path.join("UPLOADS", file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    text = extract_text(file_path, file.filename)
    if text is None:
        return {"message": "Unsupported file type"}
    chunks = chunk_text(text, file.filename)
    index_chunks(chunks, file.filename)
    return {"message": "File uploaded and indexed successfully", "filename": file.filename, "chunks_indexed": len(chunks)}


@app.post("/export/docx")
def export_docx(request: ExportRequest):
    os.makedirs("EXPORTS", exist_ok=True)
    doc = Document()
    doc.add_heading("AI Knowledge Assistant Export", level=1)
    doc.add_paragraph(f"Question: {request.question}")
    doc.add_paragraph(f"Answer: {request.answer}")
    doc.add_heading("Sources", level=2)
    for source in request.sources:
        doc.add_paragraph(str(source))
    filename = "EXPORTS/answer.docx"
    doc.save(filename)
    return {"message": "Word file exported successfully", "file": filename}


@app.post("/export/xlsx")
def export_xlsx(request: ExportRequest):
    os.makedirs("EXPORTS", exist_ok=True)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "AI Answers"
    sheet.append(["Question", "Answer", "Sources"])
    source_text = ", ".join(str(source) for source in request.sources)
    sheet.append([request.question, request.answer, source_text])
    filename = "EXPORTS/answer.xlsx"
    workbook.save(filename)
    return {"message": "Excel file exported successfully", "file": filename}


def extract_text(file_path: str, filename: str):

    if filename.lower().endswith(".pdf"):

        reader = PdfReader(file_path)

        return " ".join(
            page.extract_text()
            for page in reader.pages
            if page.extract_text()
        )

    elif filename.lower().endswith(".docx"):

        return docx2txt.process(file_path)

    elif filename.lower().endswith(".csv"):

        df = pd.read_csv(file_path)

        return df.to_string(index=False)

    elif filename.lower().endswith(".md"):

        with open(
            file_path,
            "r",
            encoding="utf-8"
        ) as f:

            return f.read()

    return None


def chunk_text(text: str, filename: str, chunk_size: int = 30):
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size):
        chunk = " ".join(words[i:i + chunk_size])
        chunks.append({"id": f"{filename}_{i}", "text": chunk})
    return chunks


def index_chunks(chunks: list, filename: str):
    client = chromadb.PersistentClient(path="./vector_db")
    collection = client.get_or_create_collection("knowledge_base")
    for chunk in chunks:
        embedding = embedding_model.encode(chunk["text"]).tolist()
        collection.add(
            ids=[chunk["id"]],
            documents=[chunk["text"]],
            embeddings=[embedding],
            metadatas=[{"source": filename}],
        )
